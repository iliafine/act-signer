"""
Поиск блока подписантов в .docx акте и вставка изображений подписей.

Приоритет сопоставления — по ФИО: в реальных актах рядом с местом для
подписи обычно напечатана фамилия с инициалами ("Иванов И.И."). Это
надёжнее, чем сопоставление по тексту должности, поэтому роль/должность
используется только как запасной вариант (и только если она задана
у подписанта в базе).

Подписанты в актах встречаются в двух типичных видах:
  1) Таблица: одна ячейка — текст должности/роли, соседняя ячейка —
     место для подписи (пустое, с подчёркиванием, и/или с напечатанным ФИО).
  2) Параграфы: строка с должностью, далее строка с подчёркиванием и/или ФИО.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import numpy as np
import docx
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.table import _Cell
from lxml import etree
from PIL import Image

from app.signature_db import Signer, find_signer_by_name, find_signer_for_role, load_db

MIN_ROLE_LEN = 12
UNDERSCORE_CHARS = set("_—–-.")
NAME_PATTERN = re.compile(r"[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.\s?[А-ЯЁ]\.|[А-ЯЁ]\.\s?[А-ЯЁ]\.\s?[А-ЯЁ][а-яё]+")

CAPTION_KEYWORDS = ("подпис",)  # "(фамилия, инициалы, подпись)" и похожие варианты


@dataclass
class SignerSlot:
    role_text: str
    kind: str  # "table_cell" | "paragraph"
    name_text: str | None = None
    table_idx: int | None = None
    row_idx: int | None = None
    role_cell_idx: int | None = None
    target_cell_idx: int | None = None
    paragraph_idx: int | None = None
    target_paragraph_idx: int | None = None


@dataclass
class SignResult:
    role_text: str
    status: str  # "signed" | "not_found_in_db" | "ambiguous" | "no_target_location"
    matched_signer: Signer | None = None
    match_score: int = 0
    matched_by: str | None = None  # "name" | "role"


@dataclass
class SignReport:
    slots: list[SignerSlot] = field(default_factory=list)
    results: list[SignResult] = field(default_factory=list)

    @property
    def all_signed(self) -> bool:
        return len(self.results) > 0 and all(r.status == "signed" for r in self.results)

    @property
    def missing(self) -> list[SignResult]:
        return [r for r in self.results if r.status != "signed"]


def _is_role_like(text: str) -> bool:
    text = text.strip()
    if len(text) < MIN_ROLE_LEN:
        return False
    if all(ch in UNDERSCORE_CHARS or ch.isspace() for ch in text):
        return False
    letters = sum(ch.isalpha() for ch in text)
    return letters / max(len(text), 1) > 0.5


def _is_blank_or_underscore(text: str) -> bool:
    text = text.strip()
    if not text:
        return True
    return all(ch in UNDERSCORE_CHARS or ch.isspace() for ch in text)


def _extract_name(text: str) -> str | None:
    match = NAME_PATTERN.search(text)
    return match.group(0) if match else None


def _is_pure_name(text: str) -> bool:
    """True, если текст ячейки/параграфа — это, по сути, только ФИО
    (не описание должности). Такие ячейки не должны сами по себе
    считаться отдельным слотом подписанта."""
    name = _extract_name(text)
    if not name:
        return False
    remainder = text.strip().replace(name, "", 1).strip(" ,.;:()")
    return len(remainder) <= 3


def _is_signature_target(text: str) -> bool:
    text = text.strip()
    if _is_blank_or_underscore(text):
        return True
    if _extract_name(text):
        return True
    return len(text) <= 60 and not text.endswith((".", ";")) and sum(ch.isalpha() for ch in text) / max(len(text), 1) < 0.9


def find_signer_slots(document: DocumentType) -> list[SignerSlot]:
    slots: list[SignerSlot] = []

    for t_idx, table in enumerate(document.tables):
        for r_idx, row in enumerate(table.rows):
            cells = row.cells
            for c_idx, cell in enumerate(cells):
                text = cell.text.strip()
                if not _is_role_like(text) or _is_pure_name(text):
                    continue
                target_idx = None
                for other_idx in (c_idx + 1, c_idx - 1):
                    if 0 <= other_idx < len(cells) and other_idx != c_idx:
                        if _is_signature_target(cells[other_idx].text):
                            target_idx = other_idx
                            break
                name_text = _extract_name(cells[target_idx].text) if target_idx is not None else None
                slots.append(
                    SignerSlot(
                        role_text=text,
                        kind="table_cell",
                        name_text=name_text,
                        table_idx=t_idx,
                        row_idx=r_idx,
                        role_cell_idx=c_idx,
                        target_cell_idx=target_idx,
                    )
                )

    paragraph_slots = _find_paragraph_slots_by_caption(document)
    if not paragraph_slots:
        paragraph_slots = _find_paragraph_slots_by_heuristic(document)
    slots.extend(paragraph_slots)

    return slots


def _find_paragraph_slots_by_caption(document: DocumentType) -> list[SignerSlot]:
    """Основной способ для параграфных актов: реальные акты почти всегда
    содержат подпись "(фамилия, инициалы, подпись)" сразу под местом для
    подписи — это надёжный якорь. От него поднимаемся вверх: сначала ищем
    строку с ФИО, затем — строку с описанием должности (пропуская пустые
    параграфы между ними)."""
    paragraphs = document.paragraphs
    slots: list[SignerSlot] = []

    for idx, para in enumerate(paragraphs):
        text_low = para.text.strip().lower()
        if not any(kw in text_low for kw in CAPTION_KEYWORDS):
            continue
        if "фамилия" not in text_low and "инициал" not in text_low:
            continue

        name_idx = idx - 1
        while name_idx >= 0 and not paragraphs[name_idx].text.strip():
            name_idx -= 1
        if name_idx < 0:
            continue

        role_idx = name_idx - 1
        while role_idx >= 0 and not paragraphs[role_idx].text.strip():
            role_idx -= 1
        if role_idx < 0:
            continue

        role_text = paragraphs[role_idx].text.strip()
        name_para_text = paragraphs[name_idx].text.strip()
        if not _is_role_like(role_text):
            continue

        slots.append(
            SignerSlot(
                role_text=role_text,
                kind="paragraph",
                name_text=_extract_name(name_para_text) or name_para_text,
                paragraph_idx=role_idx,
                target_paragraph_idx=name_idx,
            )
        )

    return slots


def _find_paragraph_slots_by_heuristic(document: DocumentType) -> list[SignerSlot]:
    """Запасной способ, если в документе нет стандартной подписи
    "(фамилия, инициалы, подпись)". Менее надёжен: рассматривает только
    нижнюю половину документа и короткие строки-кандидаты."""
    paragraphs = document.paragraphs
    start_idx = len(paragraphs) // 2
    slots: list[SignerSlot] = []

    for p_idx, para in enumerate(paragraphs):
        if p_idx < start_idx:
            continue
        text = para.text.strip()
        if not _is_role_like(text) or _is_pure_name(text):
            continue
        if len(text) > 220 or text.endswith((".", ",", ";")):
            continue

        target_p_idx = None
        name_text = None
        if p_idx + 1 < len(paragraphs):
            next_text = paragraphs[p_idx + 1].text
            if _is_signature_target(next_text):
                target_p_idx = p_idx + 1
                name_text = _extract_name(next_text)

        slots.append(
            SignerSlot(
                role_text=text,
                kind="paragraph",
                name_text=name_text,
                paragraph_idx=p_idx,
                target_paragraph_idx=target_p_idx,
            )
        )

    return slots


def _make_floating(run, offset_x_emu: int = 0, offset_y_emu: int = 0) -> None:
    """Превращает обычную (inline) картинку в "плавающую": она будет
    нарисована поверх текста (behindDoc=0) без сдвига соседних строк/ячеек —
    как если бы документ подписали от руки на печатном месте для подписи."""
    drawing = run._element.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    docPr = inline.find(qn("wp:docPr"))
    graphic = inline.find(qn("a:graphic"))

    cx, cy = extent.get("cx"), extent.get("cy")
    doc_id, doc_name = docPr.get("id"), docPr.get("name")

    inline.remove(graphic)

    anchor_xml = (
        '<wp:anchor xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" '
        'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="column"><wp:posOffset>{offset_x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>{offset_y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        "<wp:wrapNone/>"
        f'<wp:docPr id="{doc_id}" name="{doc_name}"/>'
        "</wp:anchor>"
    )
    anchor = etree.fromstring(anchor_xml)
    anchor.append(graphic)
    drawing.remove(inline)
    drawing.append(anchor)


# Габариты бокса, в который вписывается подпись. Масштабируем с сохранением
# пропорций так, чтобы подпись была соразмерна тексту (~12pt ≈ 0.42 см) и не
# выглядела громоздкой. Ограничение по ВЫСОТЕ — ключевое: иначе "высокие"
# росчерки (где ширина << высоты) при масштабе по ширине раздувались по высоте.
SIG_MAX_W_CM = 4.0
SIG_MAX_H_CM = 0.95

# Линия ("полка") под строкой с ФИО находится примерно на этом расстоянии
# ниже верха параграфа с фамилией. Низ подписи сажаем чуть ниже линии
# (небольшой "хвост" под линию — как у настоящей подписи).
SIG_LINE_OFFSET_CM = 0.42
SIG_BELOW_LINE_CM = 0.08


def _fit_box(image_path: str, scale: float = 1.0) -> tuple[float, float]:
    """Возвращает (width_cm, height_cm), вписывая картинку в бокс
    (SIG_MAX_W_CM × SIG_MAX_H_CM) × scale с сохранением пропорций.
    scale > 1 — для размашистых подписей, которые должны выходить за линию."""
    with Image.open(image_path) as im:
        w_px, h_px = im.size
    aspect = w_px / max(h_px, 1)
    max_w, max_h = SIG_MAX_W_CM * scale, SIG_MAX_H_CM * scale
    width = max_h * aspect
    if width <= max_w:
        return width, max_h
    return max_w, max_w / aspect


def _baseline_frac(image_path: str) -> float:
    """Доля высоты (от верха), на которой находится "строка" подписи —
    нижний край основной массы букв (а не габаритный низ). Нужно, чтобы
    сажать на полку именно буквы, а свисающие петли/хвосты пусть уходят
    под линию. Берём самую нижнюю строку пикселей, где плотность чернил
    ещё заметна относительно пиковой."""
    with Image.open(image_path) as im:
        alpha = np.array(im.convert("RGBA"))[:, :, 3]
    rows = (alpha > 40).sum(axis=1).astype(float)
    if rows.max() <= 0:
        return 1.0
    thr = 0.45 * rows.max()
    idx = np.where(rows >= thr)[0]
    return (idx.max() + 1) / len(rows) if len(idx) else 1.0


def _name_is_right_aligned(para) -> bool:
    """ФИО может быть выровнено по правому краю (тогда подпись ставится
    слева от фамилии). Определяем по выравниванию параграфа или по табу,
    уводящему ФИО вправо."""
    if para.alignment is not None and int(para.alignment) in (2, 3):  # RIGHT, JUSTIFY
        return True
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        tabs = pPr.find(qn("w:tabs"))
        if tabs is not None:
            for tab in tabs.findall(qn("w:tab")):
                if tab.get(qn("w:val")) == "right" and "\t" in para.text:
                    return True
    return False


def _insert_image_in_cell(cell: _Cell, image_path: str, scale: float = 1.0) -> None:
    width_cm, _ = _fit_box(image_path, scale)
    paragraph = cell.paragraphs[0] if cell.paragraphs and not cell.paragraphs[0].text.strip() else cell.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    _make_floating(run, offset_x_emu=Cm(0.3).emu, offset_y_emu=Cm(-0.25).emu)


def _insert_floating_image_in_paragraph(
    document: DocumentType, paragraph_idx: int, image_path: str, scale: float = 1.0
) -> None:
    """Вставляет подпись плавающей картинкой в строку с ФИО, сидя на линии
    ("полке"). Если ФИО слева — подпись справа от него; если ФИО выровнено
    по правому краю — подпись слева. Низ росчерка садится на линию."""
    width_cm, height_cm = _fit_box(image_path, scale)
    para = document.paragraphs[paragraph_idx]
    run = para.add_run()
    run.add_picture(image_path, width=Cm(width_cm))

    if _name_is_right_aligned(para):
        # ФИО у правого края → подпись слева от него (по центру-слева полки).
        offset_x = Cm(5.5)
    else:
        # ФИО слева → подпись справа от фамилии.
        offset_x = Cm(4.3)

    # Y: габаритный низ подписи садится на полку (с небольшим хвостом вниз).
    # Этот вариант проверен — при нём Рожин/Кудлай ложатся ровно на линию.
    bottom_cm = SIG_LINE_OFFSET_CM + SIG_BELOW_LINE_CM
    _make_floating(run, offset_x_emu=offset_x.emu, offset_y_emu=Cm(bottom_cm - height_cm).emu)


def _resolve_signer(slot: SignerSlot, signers_db: list[Signer]) -> tuple[Signer | None, int, str | None]:
    """Приоритет: совпадение по ФИО, затем (запасной вариант) по роли."""
    if slot.name_text:
        signer, score = find_signer_by_name(slot.name_text, signers_db)
        if signer is not None:
            return signer, score, "name"

    signer, score = find_signer_for_role(slot.role_text, signers_db)
    if signer is not None:
        return signer, score, "role"

    return None, score, None


def sign_document(input_path: str, output_path: str, signers_db: list[Signer] | None = None) -> SignReport:
    signers_db = signers_db if signers_db is not None else load_db()
    document = docx.Document(input_path)
    slots = find_signer_slots(document)
    report = SignReport(slots=slots)

    for slot in slots:
        matched, score, matched_by = _resolve_signer(slot, signers_db)
        if matched is None:
            report.results.append(SignResult(role_text=slot.role_text, status="not_found_in_db", match_score=score))
            continue

        sig_path = matched.signature_path()
        if not sig_path.exists():
            report.results.append(
                SignResult(role_text=slot.role_text, status="not_found_in_db", matched_signer=matched, match_score=score, matched_by=matched_by)
            )
            continue

        if slot.kind == "table_cell":
            if slot.target_cell_idx is None:
                report.results.append(
                    SignResult(role_text=slot.role_text, status="no_target_location", matched_signer=matched, match_score=score, matched_by=matched_by)
                )
                continue
            table = document.tables[slot.table_idx]
            cell = table.rows[slot.row_idx].cells[slot.target_cell_idx]
            _insert_image_in_cell(cell, str(sig_path), scale=matched.scale)
        else:
            insert_after = slot.target_paragraph_idx if slot.target_paragraph_idx is not None else slot.paragraph_idx
            _insert_floating_image_in_paragraph(document, insert_after, str(sig_path), scale=matched.scale)

        report.results.append(
            SignResult(role_text=slot.role_text, status="signed", matched_signer=matched, match_score=score, matched_by=matched_by)
        )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return report
